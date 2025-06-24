#!/usr/bin/env python3
"""
Enhanced Folder Scanner Script

This script scans a specified folder (and all its nested subfolders) for all files.
It outputs a directory tree structure and then detailed information for each file,
including filename, file type, location, and attempts to extract content and titles
for supported file types (currently .txt and .docx).

The output is saved to a text file named "folder_content_report.txt" in the
same directory as this script.

**REQUIREMENTS:**
This script requires the 'python-docx' library to process .docx files.
Install it using pip:
    pip install python-docx

Usage:
    python enhanced_folder_scanner.py [/path/to/folder]
    If no path is provided, the script will use the DEFAULT_FOLDER_PATH defined below.
"""

import os
import sys
import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("--------------------------------------------------------------------")
    print("WARNING: The 'python-docx' library is not installed or not found.")
    print("Microsoft Word (.docx) file content and title extraction will be SKIPPED.")
    print("To enable .docx processing, please install the library by running:")
    print("  pip install python-docx")
    print("--------------------------------------------------------------------")
    Document = None # Placeholder if import fails

#-------------------------------------------------------------------------
# CONFIGURATION - CHANGE THIS SECTION AS NEEDED
#-------------------------------------------------------------------------
# Set the path to the folder you want to scan here
# IMPORTANT: Use forward slashes (/) or double backslashes (\\) for Windows paths
DEFAULT_FOLDER_PATH = r"G:\My Drive\A_Capstone Thesis - Sintica\Docs\Exp_Dev_Env\ITP Engine"  # <- CHANGE THIS PATH if not using CLI arg

# --- NEW ---
# Add any folder names you want to completely exclude from the scan.
# This is case-sensitive. For example, 'backups' will be excluded but 'Backups' will not.
EXCLUDE_FOLDERS = [
    'to_exclude',
    '.git',
    '__pycache__',
    'node_modules',
    'backups',
    'temp',
    'libs',
]

# Name of the output file (will be created in the same directory as this script)
OUTPUT_FILE = "folder_content_report.txt"

# Content preview settings
MAX_CONTENT_PREVIEW_LINES = 50
MAX_CONTENT_PREVIEW_CHARS = 2000
#-------------------------------------------------------------------------

def get_file_title_and_content(filepath):
    """
    Attempts to extract a title and content from a given file.
    """
    path_obj = Path(filepath)
    extension = path_obj.suffix.lower()
    title = None
    content_preview = ""
    content_notes = ""

    try:
        if extension == ".txt":
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
            if lines:
                for line in lines: # Heuristic for title: first non-empty line
                    if line.strip():
                        title = line.strip()
                        if len(title) > 200: title = title[:200] + "..." # Cap title length
                        break
                full_content = "".join(lines)
                if len(full_content) > MAX_CONTENT_PREVIEW_CHARS:
                    content_preview = full_content[:MAX_CONTENT_PREVIEW_CHARS] + "\n... (content truncated)"
                elif len(lines) > MAX_CONTENT_PREVIEW_LINES:
                    content_preview = "".join(lines[:MAX_CONTENT_PREVIEW_LINES]) + "\n... (content truncated)"
                else:
                    content_preview = full_content
            else:
                content_notes = "[Empty text file]"

        elif extension == ".docx" and Document is not None:
            doc = Document(filepath)
            titles_found = []
            for para in doc.paragraphs:
                if para.style and para.style.name and para.style.name.lower().startswith('heading'):
                    if para.text.strip():
                        titles_found.append(para.text.strip())
            
            if titles_found:
                title = "; ".join(titles_found[:3])
                if len(title) > 200: title = title[:200] + "..."
            elif doc.paragraphs:
                for para in doc.paragraphs:
                    if para.text.strip():
                        title_candidate = para.text.strip()
                        title = title_candidate[:150] + "..." if len(title_candidate) > 150 else title_candidate
                        break

            full_text_list = [para.text for para in doc.paragraphs]
            full_text = "\n".join(full_text_list)

            if len(full_text) > MAX_CONTENT_PREVIEW_CHARS:
                content_preview = full_text[:MAX_CONTENT_PREVIEW_CHARS] + "\n... (content truncated)"
            elif len(full_text.splitlines()) > MAX_CONTENT_PREVIEW_LINES:
                content_preview = "\n".join(full_text.splitlines()[:MAX_CONTENT_PREVIEW_LINES]) + "\n... (content truncated)"
            else:
                content_preview = full_text
            if not content_preview.strip() and not title:
                content_notes = "[DOCX appears empty or has no extractable text/headings]"
            elif not content_preview.strip() and title:
                content_notes = "[DOCX has headings but no other significant body text found for preview]"

        elif extension in ['.pdf', '.xlsx', '.xls', '.ppt', '.pptx']:
            content_notes = f"[Content extraction for {extension} not yet implemented, but file is present.]"
            
        elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.svg', '.exe', '.dll', '.app', '.bin', '.zip', '.gz', '.tar', '.rar', '.7z', '.mp3', '.wav', '.aac', '.flac', '.mp4', '.avi', '.mov', '.mkv', '.webm', '.lnk', '.url']:
            content_notes = f"[Binary, media, archive, or shortcut file ({extension}). Content not displayed.]"
        else:
            try:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    lines = f.readlines()
                if lines:
                    full_content = "".join(lines)
                    if len(full_content) > MAX_CONTENT_PREVIEW_CHARS:
                        content_preview = full_content[:MAX_CONTENT_PREVIEW_CHARS] + "\n... (content truncated)"
                    elif len(lines) > MAX_CONTENT_PREVIEW_LINES:
                        content_preview = "".join(lines[:MAX_CONTENT_PREVIEW_LINES]) + "\n... (content truncated)"
                    else:
                        content_preview = full_content
                    if content_preview.strip():
                        content_notes = f"[Attempted text extraction for unknown type {extension}]"
                    else:
                        content_notes = f"[Unknown file type ({extension}), appears empty or unreadable as text]"
                else:
                    content_notes = f"[Unknown file type ({extension}), appears empty or unreadable as text]"
            except Exception:
                content_notes = f"[Unknown file type ({extension}), likely binary or not text-readable]"

    except Exception as e:
        content_notes = f"[ERROR processing file {path_obj.name}: {type(e).__name__} - {e}]"
        content_preview = ""
        title = None

    return title, content_preview, content_notes

# --- MODIFIED: REWRITTEN FOR CORRECT HIERARCHY ---
def _tree_generator(directory: Path, file_out, prefix: str = ''):
    """Recursive helper function to generate a proper directory tree."""
    # Get all items in the directory, filter out excluded folders, and sort them
    try:
        items = sorted([item for item in directory.iterdir() if item.is_dir() and item.name not in EXCLUDE_FOLDERS] +
                       [item for item in directory.iterdir() if item.is_file()])
    except PermissionError:
        file_out.write(f"{prefix}└── [Error: Permission Denied]\n")
        return

    # Define connectors for the tree branches
    branch = '├── '
    last_branch = '└── '
    
    for i, item in enumerate(items):
        connector = last_branch if i == len(items) - 1 else branch
        
        # Write the current item to the file
        if item.is_dir():
            file_out.write(f"{prefix}{connector}{item.name}/\n")
            # Determine the prefix for the next level of recursion
            new_prefix = prefix + ('    ' if i == len(items) - 1 else '│   ')
            _tree_generator(item, file_out, prefix=new_prefix)
        else:
            file_out.write(f"{prefix}{connector}{item.name}\n")


def generate_directory_tree(folder_path_str, output_file_object):
    """
    Generates a directory tree structure and writes it to the file,
    respecting the EXCLUDE_FOLDERS list.
    """
    base_folder_path = Path(folder_path_str)
    output_file_object.write(f"Directory Tree for: {base_folder_path}\n")
    output_file_object.write("="*50 + "\n")
    
    # Start the tree with the root folder name
    output_file_object.write(f"{base_folder_path.name}/\n")
    # Call the recursive generator to build the tree
    _tree_generator(base_folder_path, output_file_object)
    
    output_file_object.write("="*50 + "\n\n")
# --- END MODIFIED SECTION ---


def scan_folder_and_collect_files(folder_path):
    """
    Recursively scan a folder and collect all file paths,
    respecting the EXCLUDE_FOLDERS list.
    """
    all_files_paths = []
    root_path_obj = Path(folder_path)

    if not root_path_obj.exists():
        print(f"Error: The path '{folder_path}' does not exist.")
        sys.exit(1)
    if not root_path_obj.is_dir():
        print(f"Error: '{folder_path}' is not a directory.")
        sys.exit(1)

    for root, dirs, files in os.walk(folder_path):
        # Exclude specified directories from traversal.
        dirs[:] = [d for d in dirs if d not in EXCLUDE_FOLDERS]

        for file_name in files:
            all_files_paths.append(str(Path(root) / file_name))
    
    all_files_paths.sort()
    return all_files_paths


def write_report_to_file(files_paths, output_file_path_obj, source_folder_str):
    """
    Write the collected file information, titles, and content to the output file.
    """
    with open(output_file_path_obj, 'w', encoding='utf-8') as f:
        f.write(f"FOLDER CONTENT REPORT\n")
        f.write(f"Source Folder: {source_folder_str}\n")
        f.write(f"Scan Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Total files processed: {len(files_paths)}\n")
        f.write(f"(Excluded folders: {', '.join(EXCLUDE_FOLDERS) if EXCLUDE_FOLDERS else 'None'})\n")
        f.write("="*80 + "\n\n")

        print("Generating directory tree...")
        generate_directory_tree(source_folder_str, f)
        
        f.write("DETAILED FILE INFORMATION:\n")
        f.write("="*80 + "\n\n")

        for i, file_path_str_item in enumerate(files_paths, 1):
            file_path_item = Path(file_path_str_item)
            print(f"Processing file {i}/{len(files_paths)}: {file_path_item.name}")

            f.write(f"--- File #{i} ---\n")
            f.write(f"Filename: {file_path_item.name}\n")
            
            file_extension = file_path_item.suffix if file_path_item.suffix else "[no extension]"
            f.write(f"Type: {file_extension.lower() if file_extension != '[no extension]' else file_extension}\n")
            
            f.write(f"Location: {str(file_path_item)}\n")

            title, content_preview, content_notes = get_file_title_and_content(str(file_path_item))

            if title:
                f.write(f"Extracted Title(s)/Heading(s): {title}\n")
            
            if content_notes:
                f.write(f"Notes: {content_notes}\n")

            if content_preview and content_preview.strip():
                f.write("Content Preview:\n\"\"\"\n")
                f.write(content_preview.strip())
                f.write("\n\"\"\"\n")
            elif not content_notes or ("[Empty text file]" not in content_notes and "appears empty" not in content_notes and "Binary" not in content_notes and "not yet implemented" not in content_notes):
                f.write("Content Preview: [Not available, file might be empty, or an issue occurred during extraction]\n")

            f.write("\n" + "-"*60 + "\n\n")
    
    print(f"\nScan complete! Results written to {output_file_path_obj}")
    print(f"Processed {len(files_paths)} files in total.")


def main():
    """Main function to execute the script."""
    if len(sys.argv) > 1:
        folder_path_arg = sys.argv[1]
    else:
        folder_path_arg = DEFAULT_FOLDER_PATH
        
        if not Path(DEFAULT_FOLDER_PATH).exists():
            print(f"WARNING: The default path '{DEFAULT_FOLDER_PATH}' does not exist.")
            print("Please change the DEFAULT_FOLDER_PATH variable in the script to your target folder,")
            print("or provide a valid path as a command-line argument.")
            sys.exit("Script cannot proceed with a non-existent default path.")

    try:
        resolved_folder_path = str(Path(folder_path_arg).resolve(strict=True))
    except FileNotFoundError:
        print(f"ERROR: The specified folder path does not exist: {folder_path_arg}")
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: Invalid folder path specified: {folder_path_arg} ({e})")
        sys.exit(1)

    print(f"Starting scan for folder: {resolved_folder_path}")
    if EXCLUDE_FOLDERS:
        print(f"Excluding folders named: {', '.join(EXCLUDE_FOLDERS)}")
    
    collected_files = scan_folder_and_collect_files(resolved_folder_path)
    
    script_dir = Path(__file__).parent.resolve()
    output_path = script_dir / OUTPUT_FILE
    
    if not collected_files:
        print("No files found in the specified directory (after exclusions).")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"FOLDER CONTENT REPORT\n")
            f.write(f"Source Folder: {resolved_folder_path}\n")
            f.write(f"Scan Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("No files found in the specified directory (after exclusions).\n")
        print(f"Empty report written to {output_path}")
        return

    write_report_to_file(collected_files, output_path, resolved_folder_path)

if __name__ == "__main__":
    main()